# -*- coding: utf-8 -*-
"""
가톨릭대 공지 에이전트 (URL 재사용 + 검색 백업)
- searchURL.py가 만든 outputs/notice_links.csv의 normalized_url을 우선 사용
- 없으면 제목 검색 → 상세 링크(articleNo) 추출(정규화) → 본문 수집 → 영어 번역 → DOCX 저장
- 텍스트 중심: 이미지/표 제외

필요:
  pip install playwright python-docx python-dotenv google-generativeai
  playwright install
  .env: GEMINI_API_KEY=..., GEMINI_MODEL=gemini-1.5-pro (옵션)
실행:
  python cuk_notice_agent.py
"""

from __future__ import annotations
import os, re, time, json, csv
from dataclasses import dataclass
from typing import Optional, List, Tuple, Dict
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
import google.generativeai as genai

# --------------- 설정 ---------------
NOTICE_LIST_URL = "https://www.catholic.ac.kr/ko/campuslife/notice.do"
HEADLESS = True
BETWEEN_QUERIES_SLEEP = 0.8
NETWORK_IDLE_WAIT_MS = 1000
CSV_PATH = Path("outputs/notice_links.csv")   # searchURL.py 결과 파일

# CSV가 없을 때만 사용될 제목 리스트(예시)
TARGET_TITLES: List[str] = [
    "[건축팀] 김수환관 승강기 4호기(전망용 승강기 중 좌측) 제어시스템 교체공사 안내",
    "[대외협력팀] 홍보 콘텐츠 제작 대외협력팀 소속 기관동아리 <CUK프렌즈> 18기 모집 공고"
]

TRANSLATION_PROMPT_TEMPLATE = """You are a professional translator for university announcements.
Translate Korean to clear, natural, formal English.

Rules for Translation:
- Text-only output. Keep lists and section headings tidy.
- Translate 'Catholic University of Korea' to 'The Catholic University of Korea'.
- For building and department names, use the following official English names:
  * 정문: Main Gate
  * 창업보육센터(BI관): Business Incubation Center(BI)
  * 김수환관(K관): Kim Sou-hwan Hall(K)
  * 스테파노 기숙사(K관): Dormitory Stephen(K)
  * 안드레아관(A관): Andreas Dormitory(A)
  * 마리아관(M관): Maria Hall(M)
  * 니콜스관(N관): Nicholls Hall(N)
  * 밤비노관(BA관): Bambino Hall(BA)
  * 다솔관(D관): Dasol Hall (Science Hall)(D)
  * 비르투스관(V관): Virtus Hall(V)
  * 학생미래인재관(B관): Sophie Barat Hall(B)
  * 하늘동산: Outdoor Stage
  * 미카엘관-행정동(H관): Michael Hall (Administration)(H)
  * 미카엘관-교수연구동(T관): Michael Hall (Faculty)(T)
  * 베리타스관(중앙도서관): Veritas Hall (Central Library)(L)
  * 성심관(SH관): Songsim Hall(SH)
  * 정진석추기경학부(NP관): Nicholas Cardinal Cheong Pharmacy Hall(NP)
  * 콘서트홀(CH관): Concert Hall (Grand Auditorium)(CH)
  * 대운동장: Athletic Ground
  * 예수성심성당(C관): Chapel of the Sacred Heart(C)
  * 프란치스코 기숙사(F관): Francisco Dormitory(F)
  * 건축팀: Architecture Team
  * 교목실: Office of Campus Ministry
  * 교수학습개발원: Center for Teaching & Learning
  * 기숙사운영팀: Management Team of Dormitory
  * 대외협력팀: External Affairs Team
  * 시설팀: Facilities Management Team
  * 총무팀: General Affairs Team
  * 취창업지원팀: Career Support Team
  * 평생교육원: The Center for Lifelong Education
  * 학부대학운영팀: College Management Team
  * 학사지원팀: Academic Affairs Services Team
  * 학생지원팀: Student Affairs Team
  * 융합전공학부: School of interdisciplinary studies
  * 중앙도서관: The Catholic University of Korea Central Library
  * 인간연구소: Research Institute of Anthropology
  * 구매관재팀: Purchasing and Property Management Team
  * 창업교육혁신센터: Center for Entrepreneurship Education and Innovation
  * 일반대학원(교학팀): General Graduate School (Team)
  * 특수대학원(교학팀): Specialized Graduate School (Team)
  * 교육대학원/교회법대학원교학팀: Graduate School of Education/Graduate School of Canon Law
  * 건설본부: Construction Headquarters
  * 정보통신원: Information & Communications Center
  * 고용노동부: Ministry of Employment and Labor
  * 한국가톨릭굿뉴스: Catholic Metaversity
  * 인권센터: Human Rights Center
- Preserve bullet/numbered lists and section structures, including '△' or similar symbols. Do not rephrase them into a summary format.
- For "학보 기사" style announcements, translate '글/사진:' to 'Text/Photos:'.
- Dates: keep original format and add weekday in parentheses if present (e.g., 2025.08.18 (Mon) ~ 2025.08.22 (Fri)).
- Remove UI crumbs like '목록', '이전글/다음글', '첨부파일'.
- Do NOT add content not present in source. If unclear, write [unclear].

Original Korean text to translate:

{text_ko}
"""

# --------------- 데이터 구조 ---------------
@dataclass
class FoundNotice:
    title_ko: str
    url_raw: Optional[str]
    url_norm: Optional[str]
    status: str
    note: str = ""
    date_text: Optional[str] = None
    body_text_ko: Optional[str] = None
    body_text_en: Optional[str] = None

# --------------- 유틸 ---------------
def normalize_cuk_notice_url(url: str) -> str:
    p = urlparse(url)
    qs = parse_qs(p.query)
    article_no = qs.get("articleNo", [None])[0]
    if not article_no:
        return url
    new_qs = {"mode": "view", "articleNo": article_no}
    return urlunparse((p.scheme, p.netloc, p.path, p.params, urlencode(new_qs), ""))

def extract_article_no(url: str) -> Optional[str]:
    m = re.search(r"(?:[?&])articleNo=(\d+)", url)
    return m.group(1) if m else None

def load_urls_from_csv(path: Path) -> List[Dict[str, str]]:
    """
    notice_links.csv (columns: title,status,normalized_url,raw_url,note) 읽기
    normalized_url 우선, 없으면 raw_url 사용.
    """
    if not path.exists():
        return []
    rows = []
    with path.open("r", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            url = (row.get("normalized_url") or "").strip() or (row.get("raw_url") or "").strip()
            if not url:
                continue
            rows.append({
                "title": (row.get("title") or "").strip(),
                "url": normalize_cuk_notice_url(url)
            })
    return rows

# --------------- DOM 헬퍼 ---------------
def find_search_input(page):
    candidates = [
        "input[name='srSearchVal']",
        "input#srSearchVal",
        "input[type='search']",
        "input[placeholder*='검색']",
        "input[title*='검색']",
        "input[name*='search']",
    ]
    for sel in candidates:
        el = page.query_selector(sel)
        if el:
            return el
    els = page.query_selector_all("input")
    return els[0] if els else None

def click_first_result_and_get_url(page) -> Optional[str]:
    selectors = [
        "table tbody tr td a",        # 표 형태
        "ul li a.title",              # 리스트 형태
        ".board-list a",              # 공통 보드
        "a[href*='notice.do?mode=view']",
        "a[href*='articleNo=']",
    ]
    for sel in selectors:
        try:
            page.wait_for_selector(sel, timeout=5000)
            links = page.query_selector_all(sel)
            if links:
                with page.expect_navigation(timeout=15000):
                    links[0].click()
                return page.url
        except PWTimeout:
            continue
    return None

def extract_content_text(page) -> str:
    """
    가톨릭대 공지 상세에서 본문 텍스트 추출.
    - .fr-view (p/span 구조) 1순위 처리
    - 불필요 요소 제거 및 &nbsp; 정규화
    - 실패 시 후보 컨테이너 / 최후수단 폴백
    """
    import re as _re

    # 0) 먼저 .fr-view를 강하게 시도 (네가 준 DOM 구조)
    fr = page.query_selector(".fr-view") or page.query_selector(".b-con-box .fr-view")
    if fr:
        raw = fr.evaluate(
            """
            (node) => {
              // 불필요 요소 제거
              const clone = node.cloneNode(true);
              clone.querySelectorAll('script,style,.attach,.file,.btn,.share,.sns,.prev-next,.pagination').forEach(n => n.remove());

              const ps = Array.from(clone.querySelectorAll('p'));
              let lines = [];
              if (ps.length) {
                for (const p of ps) {
                  // p 안의 텍스트 수집 (span 등 포함)
                  let t = p.innerText || "";
                  // NBSP -> space
                  t = t.replace(/\\u00A0/g, " ");
                  // 라인 단위 불필요 공백 정리
                  t = t.replace(/[ \\t]+/g, " ").trim();
                  if (t) lines.push(t);
                }
              } else {
                let t = clone.innerText || "";
                t = t.replace(/\\u00A0/g, " ");
                t = t.replace(/[ \\t]+/g, " ").trim();
                if (t) lines.push(t);
              }
              return lines.join("\\n");
            }
            """
        ) or ""

        # 하단 UI 컷오프 및 공백 정돈
        cut_words = ["첨부파일", "이전글", "다음글", "목록", "SNS", "공유"]
        for w in cut_words:
            idx = raw.find(w)
            if idx != -1 and idx > 50:
                raw = raw[:idx].rstrip()
                break

        # 여러 빈 줄 합치기
        raw = _re.sub(r"\n{3,}", "\n\n", raw).strip()
        if len(raw) >= 30:
            return raw

    # 1) 일반 컨테이너 후보들
    body_candidates = [
        "div.view-cont", "div.viewCont", "div.view-con",
        "div.board-view .cont", "div.board-view .content", "div.board-view",
        "div.board_view", "div.bv_cont", "div.bv-cont", "div.bvCon",
        "div#content", "div#boardContent", "div.article-body", "article.post",
        "#contents .view", "#contents .board", "#container .contents",
        "#contents .post-view", "#contents .bd_view", "#contents .view-con",
    ]
    for sel in body_candidates:
        el = page.query_selector(sel)
        if not el:
            continue
        txt = (el.inner_text() or "").replace("\u00A0", " ").strip()
        if len(txt) < 30:
            continue
        for w in ["첨부파일", "이전글", "다음글", "목록", "SNS", "공유"]:
            idx = txt.find(w)
            if idx != -1 and idx > 50:
                txt = txt[:idx].rstrip()
                break
        txt = _re.sub(r"[ \t]+", " ", txt)
        txt = _re.sub(r"\n{3,}", "\n\n", txt).strip()
        if len(txt) >= 30:
            return txt

    # 2) 최후수단: 페이지에서 가장 긴 텍스트 블록
    best_txt, best_len = "", 0
    for el in page.query_selector_all("main *, #contents *, #container *")[:2000]:
        try:
            t = (el.inner_text() or "").replace("\u00A0", " ").strip()
        except Exception:
            continue
        if len(t) > best_len:
            best_len, best_txt = len(t), t
    if best_txt:
        for w in ["첨부파일", "이전글", "다음글", "목록", "SNS", "공유"]:
            idx = best_txt.find(w)
            if idx != -1 and idx > 50:
                best_txt = best_txt[:idx].rstrip()
                break
        best_txt = _re.sub(r"[ \t]+", " ", best_txt)
        best_txt = _re.sub(r"\n{3,}", "\n\n", best_txt).strip()
        return best_txt

    # 3) 마지막으로 body 전체
    full = (page.inner_text("body") or "").replace("\u00A0", " ").strip()
    full = _re.sub(r"[ \t]+", " ", full)
    full = _re.sub(r"\n{3,}", "\n\n", full).strip()
    return full


# --------------- 번역 ---------------
def translate_ko_to_en(text_ko: str) -> str:
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    model_name = os.getenv("GEMINI_MODEL", "gemini-1.5-pro")
    if not api_key:
        return f"[No GEMINI_API_KEY provided – original text follows]\n\n{text_ko}"

    genai.configure(api_key=api_key)
    try:
        model = genai.GenerativeModel(model_name=model_name)
    except Exception as e:
        return f"[Error: Failed to load model '{model_name}'. Details: {e}]\n\n{text_ko}"

    prompt = TRANSLATION_PROMPT_TEMPLATE.format(text_ko=text_ko)
    retries = 3
    for i in range(retries):
        try:
            resp = model.generate_content(prompt)
            return (resp.text or "").strip()
        except Exception as e:
            if "Quota exceeded" in str(e):
                print("Quota exceeded. Waiting for 60 seconds...")
                time.sleep(60)
                continue
            if i < retries - 1:
                print(f"Error calling Gemini API. Retrying in 2 seconds... ({e})")
                time.sleep(2)
            else:
                return f"[Error translating text after {retries} retries: {e}]\n\n{text_ko}"

# --------------- 파이프라인 ---------------
def fetch_from_url(page, title: str, url: str) -> FoundNotice:
    """정규화 URL을 직접 열어 본문 추출/번역"""
    item = FoundNotice(title_ko=title or "(no title)", url_raw=url, url_norm=normalize_cuk_notice_url(url), status="pending")
    try:
        print(f"\n--- Processing URL: {item.url_norm} ---")
        page.goto(item.url_norm, wait_until="domcontentloaded", timeout=20000)
        try:
            page.wait_for_load_state("networkidle", timeout=8000)
        except PWTimeout:
            pass
        page.wait_for_timeout(300)

        body_ko = extract_content_text(page)
        item.body_text_ko = body_ko or ""

        parts = [title] if title else []
        if item.body_text_ko:
            parts.append(item.body_text_ko)
        ko_bundle = "\n\n".join(parts).strip() or "(empty)"
        
        # 크롤링된 한국어 원문 출력
        print("\n📄 Korean Original Text (bundle):")
        print("--------------------------------")
        print(ko_bundle)

        item.body_text_en = translate_ko_to_en(ko_bundle)
        item.status = "ok"
        item.note = ""

        # 번역된 영어 문장 출력
        print("\n English Translation:")
        print("----------------------")
        print(item.body_text_en)

    except Exception as e:
        item.status = "error"
        item.note = f"{type(e).__name__}: {e}"
    return item

def fetch_by_search(page, title: str) -> FoundNotice:
    """CSV가 없을 때: 제목 검색으로 링크 획득 후 본문/번역"""
    item = FoundNotice(title_ko=title, url_raw=None, url_norm=None, status="pending")
    try:
        print(f"\n--- Searching for: {title} ---")
        page.goto(NOTICE_LIST_URL, wait_until="domcontentloaded", timeout=20000)
        page.wait_for_timeout(400)
        si = find_search_input(page)
        if not si:
            item.status = "fail"; item.note = "검색창 탐색 실패"; return item

        si.fill(""); si.type(title); si.press("Enter")
        try:
            page.wait_for_load_state("networkidle", timeout=10000)
        except PWTimeout:
            pass
        page.wait_for_timeout(NETWORK_IDLE_WAIT_MS)

        detail_url = click_first_result_and_get_url(page)
        if not detail_url:
            item.status = "not_found"; item.note = "상세 링크 미발견"; return item

        item.url_raw = detail_url
        item.url_norm = normalize_cuk_notice_url(detail_url)

        try:
            page.wait_for_load_state("networkidle", timeout=8000)
        except PWTimeout:
            pass
        page.wait_for_timeout(300)

        body_ko = extract_content_text(page)
        item.body_text_ko = body_ko or ""

        parts = [title]
        if item.body_text_ko:
            parts.append(item.body_text_ko)
        ko_bundle = "\n\n".join(parts).strip()
        
        # 크롤링된 한국어 원문 출력
        print("\n📄 Korean Original Text (bundle):")
        print("--------------------------------")
        print(ko_bundle)

        item.body_text_en = translate_ko_to_en(ko_bundle)
        item.status = "ok"
        item.note = ""
        
        # 번역된 영어 문장 출력
        print("\n✅ English Translation:")
        print("----------------------")
        print(item.body_text_en)

    except Exception as e:
        item.status = "error"
        item.note = f"{type(e).__name__}: {e}"
    return item

def fetch_and_translate() -> List[FoundNotice]:
    results: List[FoundNotice] = []
    url_rows = load_urls_from_csv(CSV_PATH)  # ← 우선 CSV 사용
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=HEADLESS)
        page = browser.new_page()

        if url_rows:
            print(f"[info] Using URLs from: {CSV_PATH}")
            for r in url_rows:
                it = fetch_from_url(page, r.get("title", ""), r["url"])
                results.append(it)
                time.sleep(BETWEEN_QUERIES_SLEEP)
        else:
            print("[info] CSV not found. Falling back to title search.")
            for t in TARGET_TITLES:
                it = fetch_by_search(page, t)
                results.append(it)
                time.sleep(BETWEEN_QUERIES_SLEEP)

        browser.close()
    return results

def build_docx(items: List[FoundNotice], out_path: str):
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    h = doc.add_heading("Selected Announcements – English Compilation", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Source: Catholic University of Korea – Campus Life > Notice")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for it in items:
        sec = doc.add_heading("", level=1); sec_run = sec.add_run(it.title_ko); sec_run.bold = True

        if it.url_norm:
            doc.add_paragraph(f"Original Link: {it.url_norm}")
        elif it.url_raw:
            doc.add_paragraph(f"Original Link (raw): {it.url_raw}")

        if it.status != "ok":
            doc.add_paragraph(f"Status: {it.status} ({it.note})")
            doc.add_paragraph(""); continue

        en = it.body_text_en or ""
        for line in en.splitlines():
            doc.add_paragraph(line)
        doc.add_paragraph("\n—\n")

    for para in doc.paragraphs:
        for r in para.runs:
            r.font.size = Pt(11)

    doc.save(str(out_path))

def main():
    items = fetch_and_translate()

    Path("outputs").mkdir(exist_ok=True)
    with open("outputs/notice_results.json", "w", encoding="utf-8") as f:
        json.dump([vars(x) for x in items], f, ensure_ascii=False, indent=2)

    build_docx(items, "outputs/CUK_Announcements_EN.docx")
    print("\n Done: outputs/CUK_Announcements_EN.docx")
    for it in items:
        print(f"- [{it.status}] {it.title_ko} -> {it.url_norm or it.url_raw or 'N/A'} ({it.note})")

if __name__ == "__main__":
    main()
